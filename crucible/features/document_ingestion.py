"""
features/document_ingestion.py
================================
RAG-style document injection for the Crucible pipeline.

Reads local files (TXT, Markdown, PDF, DOCX) and produces a condensed
``ResearchContext`` supplement that can be prepended to Stage 0 / Stage 1
prompts.  Allows teams to feed existing internal research, design docs, or
prior analyses directly into the pipeline without re-typing context.

Supported formats
-----------------
* **.txt / .md / .rst** — read via stdlib (always available).
* **.pdf** — extracted via ``pypdf`` when installed; falls back to a
  base64-encoded binary notice if the library is absent.
* **.docx** — extracted via ``python-docx`` when installed; falls back to
  a plain-text notice otherwise.
* Any other extension — treated as UTF-8 text (best-effort).

Token budget
------------
Each document is truncated at ``DOCUMENT_INGESTION_MAX_CHARS`` characters
(env var, default 8 000) so that the injected prefix remains within a
reasonable context-window fraction when many documents are loaded.

The overall multi-document prefix is bounded by
``DOCUMENT_INGESTION_TOTAL_CHARS`` (env var, default 24 000).

Usage::

    from crucible.features.document_ingestion import (
        ingest_documents,
        DocumentIngestionConfig,
    )

    config = DocumentIngestionConfig(
        paths=["/docs/prior_research.pdf", "/notes/arch_decisions.md"],
    )
    result = ingest_documents(config)
    if result.context_text:
        print(result.context_text[:500])  # inject into pipeline prompt
"""
from __future__ import annotations

import os
import warnings
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

# ── Optional dependency guards ────────────────────────────────────────────────

try:
    import pypdf as _pypdf  # pypdf >= 3.x
    _HAS_PYPDF = True
except ImportError:
    try:
        import PyPDF2 as _pypdf  # type: ignore[no-redef]  # legacy fallback
        _HAS_PYPDF = True
    except ImportError:
        _pypdf = None  # type: ignore[assignment]
        _HAS_PYPDF = False

try:
    import docx as _docx  # python-docx
    _HAS_DOCX = True
except ImportError:
    _docx = None  # type: ignore[assignment]
    _HAS_DOCX = False

# ── Configuration ─────────────────────────────────────────────────────────────


try:
    from .. import _env
except ImportError:  # pragma: no cover - script-mode fallback
    import _env  # type: ignore[no-redef]


def _env_int(name: str, default: int) -> int:
    return _env.env_int(name, default)


_MAX_DOC_CHARS: int = _env_int("DOCUMENT_INGESTION_MAX_CHARS", 8_000)
_MAX_TOTAL_CHARS: int = _env_int("DOCUMENT_INGESTION_TOTAL_CHARS", 24_000)

# Extensions always treated as plain text
_TEXT_EXTENSIONS = frozenset({
    ".txt", ".md", ".markdown", ".rst", ".text",
    ".log", ".csv", ".tsv", ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".json", ".xml", ".html", ".htm",
})


# ── Data models ───────────────────────────────────────────────────────────────

@dataclass
class DocumentIngestionConfig:
    """Configuration for a single document ingestion session."""

    # File paths to ingest (absolute or relative).
    paths: List[str] = field(default_factory=list)
    # Per-document char cap (overrides env default when > 0).
    max_chars_per_doc: int = 0
    # Total char cap across all documents (overrides env default when > 0).
    max_total_chars: int = 0
    # Label prepended to the context block header.
    label: str = "Injected Document Context"


@dataclass
class IngestedDocument:
    """One successfully read document."""

    path: str
    format_detected: str          # "text" | "pdf" | "docx" | "unknown"
    char_count: int
    truncated: bool
    text: str
    error: Optional[str] = None


@dataclass
class DocumentIngestionResult:
    """Result of ingesting a batch of documents."""

    documents: List[IngestedDocument] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    context_text: str = ""        # ready-to-inject prefix string
    total_chars: int = 0

    @property
    def success_count(self) -> int:
        return sum(1 for d in self.documents if d.error is None)

    @property
    def error_count(self) -> int:
        return len(self.errors)


# ── Extraction helpers ────────────────────────────────────────────────────────

def _extract_text_file(path: str) -> Tuple[str, str]:
    """Return ``(format, text)`` for plain-text files."""
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return "text", fh.read()


def _extract_pdf(path: str) -> Tuple[str, str]:
    """Return ``(format, text)`` for PDF files."""
    if not _HAS_PYPDF or _pypdf is None:
        warnings.warn(
            "document_ingestion: pypdf / PyPDF2 not installed — cannot extract "
            f"text from '{os.path.basename(path)}'. "
            "Install with: pip install pypdf",
            stacklevel=3,
        )
        return "pdf", f"[PDF extraction unavailable — install pypdf: {path}]"

    pages_text: List[str] = []
    try:
        # Support both pypdf (3.x) PdfReader and PyPDF2 PdfReader
        reader = _pypdf.PdfReader(path)  # type: ignore[attr-defined]
        for page in reader.pages:
            try:
                pages_text.append(page.extract_text() or "")
            except Exception:
                pages_text.append("")
    except Exception as exc:
        return "pdf", f"[PDF read error: {exc}]"
    return "pdf", "\n".join(pages_text)


def _extract_docx(path: str) -> Tuple[str, str]:
    """Return ``(format, text)`` for DOCX files."""
    if not _HAS_DOCX or _docx is None:
        warnings.warn(
            "document_ingestion: python-docx not installed — cannot extract "
            f"text from '{os.path.basename(path)}'. "
            "Install with: pip install python-docx",
            stacklevel=3,
        )
        return "docx", f"[DOCX extraction unavailable — install python-docx: {path}]"

    try:
        doc = _docx.Document(path)  # type: ignore[attr-defined]
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "docx", "\n".join(paragraphs)
    except Exception as exc:
        return "docx", f"[DOCX read error: {exc}]"


def _detect_and_extract(path: str) -> Tuple[str, str, Optional[str]]:
    """
    Detect file type and extract text.

    Returns ``(format_detected, text, error_or_None)``.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in _TEXT_EXTENSIONS:
            fmt, text = _extract_text_file(path)
        elif ext == ".pdf":
            fmt, text = _extract_pdf(path)
        elif ext in (".docx", ".doc"):
            fmt, text = _extract_docx(path)
        else:
            # Attempt as UTF-8 text regardless of extension
            fmt, text = _extract_text_file(path)
            fmt = "unknown"
        return fmt, text, None
    except OSError as exc:
        return "unknown", "", str(exc)
    except Exception as exc:
        return "unknown", "", f"{type(exc).__name__}: {exc}"


# ── Main entry point ──────────────────────────────────────────────────────────

def ingest_documents(config: DocumentIngestionConfig) -> DocumentIngestionResult:
    """
    Ingest a list of local documents and return a ready-to-inject context.

    Parameters
    ----------
    config:
        ``DocumentIngestionConfig`` specifying paths and char budgets.

    Returns
    -------
    DocumentIngestionResult
        Contains per-document results, error list, and the combined
        ``context_text`` string for prompt injection.
    """
    max_doc = config.max_chars_per_doc if config.max_chars_per_doc > 0 else _MAX_DOC_CHARS
    max_total = config.max_total_chars if config.max_total_chars > 0 else _MAX_TOTAL_CHARS

    result = DocumentIngestionResult()

    for raw_path in config.paths:
        path = os.path.abspath(raw_path)
        if not os.path.isfile(path):
            msg = f"File not found: {path}"
            result.errors.append(msg)
            result.documents.append(
                IngestedDocument(
                    path=path,
                    format_detected="unknown",
                    char_count=0,
                    truncated=False,
                    text="",
                    error=msg,
                )
            )
            continue

        fmt, text, err = _detect_and_extract(path)
        if err:
            result.errors.append(f"{os.path.basename(path)}: {err}")
            result.documents.append(
                IngestedDocument(
                    path=path,
                    format_detected=fmt,
                    char_count=0,
                    truncated=False,
                    text="",
                    error=err,
                )
            )
            continue

        truncated = len(text) > max_doc
        if truncated:
            text = text[:max_doc]

        result.documents.append(
            IngestedDocument(
                path=path,
                format_detected=fmt,
                char_count=len(text),
                truncated=truncated,
                text=text,
            )
        )

    # Build combined context text within total budget
    result.context_text = _build_context_text(
        result.documents, max_total, config.label
    )
    result.total_chars = len(result.context_text)
    return result


def _build_context_text(
    documents: List[IngestedDocument],
    max_total: int,
    label: str,
) -> str:
    """Concatenate successful documents into a single prompt prefix."""
    successful = [d for d in documents if d.error is None and d.text.strip()]
    if not successful:
        return ""

    header = (
        f"=== {label} ({len(successful)} file(s)) ===\n"
        "The following documents were provided by the user as supplementary "
        "research context. Treat them as authoritative inputs when they conflict "
        "with general knowledge.\n"
    )
    _footer = "=== End Injected Context ===\n"
    budget = max(0, max_total - len(header) - len(_footer))
    parts: List[str] = [header]

    for doc in successful:
        sep = f"\n--- [{doc.format_detected.upper()}] {os.path.basename(doc.path)}"
        if doc.truncated:
            sep += " (truncated)"
        sep += " ---\n"
        block = sep + doc.text.strip() + "\n"
        if budget - len(block) < 0:
            break
        parts.append(block)
        budget -= len(block)

    if len(header) + len(_footer) <= max_total:
        parts.append(_footer)
    return "".join(parts)


# ── Standalone usage helper ───────────────────────────────────────────────────

def ingest_documents_from_dir(
    directory: str,
    *,
    extensions: Optional[List[str]] = None,
    recursive: bool = False,
    max_chars_per_doc: int = 0,
    max_total_chars: int = 0,
) -> DocumentIngestionResult:
    """
    Convenience wrapper: ingest all matching files from a directory.

    Parameters
    ----------
    directory:
        Path to scan for documents.
    extensions:
        List of extensions to include (e.g. ``[".pdf", ".md"]``).
        Defaults to all supported text-like extensions.
    recursive:
        If True, walk subdirectories recursively.
    """
    allowed = frozenset(
        ext.lower() for ext in (extensions or [
            ".txt", ".md", ".markdown", ".rst",
            ".pdf", ".docx",
        ])
    )

    paths: List[str] = []
    if recursive:
        for dirpath, _dirs, filenames in os.walk(directory):
            for fname in sorted(filenames):
                if os.path.splitext(fname)[1].lower() in allowed:
                    paths.append(os.path.join(dirpath, fname))
    else:
        try:
            for fname in sorted(os.listdir(directory)):
                fpath = os.path.join(directory, fname)
                if os.path.isfile(fpath) and os.path.splitext(fname)[1].lower() in allowed:
                    paths.append(fpath)
        except OSError:
            pass

    config = DocumentIngestionConfig(
        paths=paths,
        max_chars_per_doc=max_chars_per_doc,
        max_total_chars=max_total_chars,
    )
    return ingest_documents(config)
